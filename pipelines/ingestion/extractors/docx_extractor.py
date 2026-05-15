
import os

from docx import Document
from pipelines.ingestion.models import ExtractedPage


async def extract_docx(file_path: str):
    doc = Document(file_path)

    pages = []

    text = []

    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            text.append(" | ".join(row_text))

    pages.append(
        ExtractedPage(
            page_number=1,
            content="\n".join(text),
            content_type="text",
            metadata={"source": "docx"}
        )
    )

    return pages

